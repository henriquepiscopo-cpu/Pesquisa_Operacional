"""
Gerador do Relatório Final DEA em .docx
Formatação baseada em relatorio_metrologia_claudecomcapaprarevisar.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Caminhos ─────────────────────────────────────────────────────────────────
HERE  = os.path.dirname(os.path.abspath(__file__))
BASE  = os.path.dirname(HERE)
FIGS2 = os.path.join(BASE, "ETAPA 2", "figures")
FIGS4 = os.path.join(BASE, "ETAPA 4", "figures")
OUT   = os.path.join(HERE, "relatorio_final.docx")

# Largura útil da página em EMU: 21cm - 2×2.22cm = 16.56cm
PAGE_WIDTH = Cm(16.56)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_table_width(table, width_emu):
    """Força largura total da tabela."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(width_emu / 635)))  # EMU → twentieths of a point (twips)
    tblW.set(qn('w:type'), 'dxa')

def set_col_widths(table, widths_cm):
    """Define largura de cada coluna explicitamente."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            twips = int(Cm(widths_cm[j]) / 635)
            tcW.set(qn('w:w'), str(twips))
            tcW.set(qn('w:type'), 'dxa')

def add_run(para, text, bold=False, italic=False, size=12):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return run

def set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=6,
                 first_indent=None, left_indent=None,
                 line_spacing=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing

def section_heading(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=12, space_after=6, line_spacing=Pt(14))
    add_run(p, text, bold=True, size=12)
    return p

def subheading_inline(doc, label, text=''):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=6, space_after=6,
                 first_indent=Cm(1.25), line_spacing=Pt(14))
    add_run(p, label, bold=True, size=12)
    if text:
        add_run(p, text, size=12)
    return p

def body_para(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=6,
                 first_indent=Cm(1.25), line_spacing=Pt(14))
    add_run(p, text, size=12)
    return p

def caption_para(doc, label, text):
    """Legenda: **Figura N.** Texto — centralizado, 10pt."""
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=2, space_after=10, line_spacing=Pt(12))
    add_run(p, label, bold=True, size=10)
    add_run(p, text, size=10)
    return p

def add_figure(doc, path, fig_label, fig_desc, width=Cm(14)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=6, space_after=2, line_spacing=Pt(14))
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        p = doc.add_paragraph(f'[IMAGEM NÃO ENCONTRADA: {os.path.basename(path)}]')
        set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    caption_para(doc, fig_label, fig_desc)

def add_equation(doc, eq_text, number):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(2.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)
    # Equação à esquerda, número à direita via tab
    tab_stops = p.paragraph_format.tab_stops
    run = p.add_run(f'{eq_text}\t{number}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Tab stop para o número no final da linha
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')  # ~16.5cm em twips (1cm=566.9 twips)
    tabs.append(tab)
    pPr.append(tabs)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def make_table(doc, headers, rows, col_widths_cm, caption_label, caption_text):
    """Cria tabela com bordas, cabeçalho bold, larguras explícitas."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)

    # Tentar aplicar estilo de grade; fallback manual
    try:
        table.style = doc.styles['Table Grid']
    except KeyError:
        pass

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Definir largura total e por coluna
    set_table_width(table, PAGE_WIDTH)
    set_col_widths(table, col_widths_cm)

    def fmt_cell(cell, text, bold=False, center=True):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = Pt(12)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cabeçalho
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        fmt_cell(hdr_row.cells[i], h, bold=True)

    # Dados
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            fmt_cell(row.cells[c_i], str(val))

    # Bordas via XML para garantir que aparecem
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), '000000')
                tcBorders.append(el)
            # Remove bordas anteriores se existirem
            old = tcPr.find(qn('w:tcBorders'))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcBorders)

    doc.add_paragraph()  # espaço
    caption_para(doc, caption_label, caption_text)
    return table


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# Margens
sec = doc.sections[0]
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin   = Cm(2.22)
sec.right_margin  = Cm(2.22)

# Estilo Normal base
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(14)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after  = Pt(0)

# ════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
add_run(p, 'LE505 - Pesquisa Operacional', size=12)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
add_run(p, 'Prof.ª Priscila Rampazzo', size=12)

for _ in range(7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'Universidade Estadual de Campinas', bold=True, size=14)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'FCA – Faculdade de Ciências Aplicadas', bold=True, size=13)

for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'Análise de Eficiência Relativa de Empresas de Petróleo e Gás', size=12)
p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'por meio da Análise Envoltória de Dados (DEA)', size=12)

for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'Junho/2026', size=12)

for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)
add_run(p, 'Henrique Arbex Rodrigues Piscopo', bold=True, size=12)
add_run(p, '          RA: 284675', bold=True, size=12)

for _ in range(7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, 'Limeira', size=12)

p = doc.add_paragraph()
set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_run(p, '2026', size=12)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# RESUMO
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, 'RESUMO')
body_para(doc, (
    'O presente trabalho aplica a Análise Envoltória de Dados (DEA) para mensurar a eficiência '
    'relativa de 22 empresas do setor de petróleo e gás listadas em bolsas internacionais. '
    'Foram empregados os modelos CCR (retornos constantes de escala) e BCC (retornos variáveis '
    'de escala), ambos com orientação a output, implementados em Python com uso da biblioteca '
    'PuLP para resolução dos problemas de programação linear. Os dados financeiros foram coletados '
    'via biblioteca yfinance, referentes ao exercício fiscal mais recente disponível, e compreendem '
    'três inputs — ativos totais, despesas operacionais e número de funcionários — e dois outputs '
    '— receita líquida e EBITDA. Os resultados indicam que 9 das 22 empresas (41%) são eficientes '
    'no modelo CCR e 10 (45%) no modelo BCC, com scores médios de 1,0811 e 1,0591, respectivamente. '
    'A Equinor (EQNR) e a Petrobras (PBR) emergem como benchmarks dominantes, sendo referenciadas '
    'por 12 e 10 DMUs ineficientes no modelo CCR, respectivamente. A Canadian Natural Resources '
    '(CNQ) é a empresa mais ineficiente da amostra (\u03b8 = 1,2597), seguida da Eni (\u03b8 = 1,1942) '
    'e da Chevron (\u03b8 = 1,1849). Conclui-se que a maioria das grandes empresas integradas opera '
    'próximo à fronteira eficiente, com ineficiências moderadas atribuíveis a diferenças de modelo '
    'de negócio, escala de operação e contexto regulatório.'
))

# ════════════════════════════════════════════════════════════════════════════
# INTRODUÇÃO
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, 'INTRODUÇÃO')
body_para(doc, (
    'O setor de petróleo e gás é marcado por elevada intensidade de capital, volatilidade de '
    'preços e crescentes pressões regulatórias e ambientais. Nesse ambiente, a avaliação '
    'comparativa do desempenho das empresas torna-se instrumento estratégico para gestores, '
    'investidores e formuladores de políticas públicas. Métodos tradicionais de análise '
    'financeira, como indicadores de rentabilidade e retorno sobre ativos, avaliam empresas '
    'de forma isolada e são fortemente influenciados por variáveis de preço exógenas ao '
    'controle das firmas, limitando a capacidade de identificar ineficiências operacionais genuínas.'
))
body_para(doc, (
    'A Análise Envoltória de Dados (DEA), método não paramétrico desenvolvido por Charnes, '
    'Cooper e Rhodes (1978) e expandido por Banker, Charnes e Cooper (1984), oferece uma '
    'abordagem comparativa que constrói uma fronteira de melhores práticas a partir das próprias '
    'observações da amostra. Por meio da resolução de problemas de programação linear, o DEA '
    'identifica quais unidades de decisão (DMUs) operam na fronteira eficiente, quais estão '
    'aquém dela e em quanto cada unidade ineficiente precisaria expandir seus outputs para '
    'tornar-se eficiente.'
))
body_para(doc, (
    'Este trabalho tem por objetivo aplicar os modelos DEA-CCR e DEA-BCC, com orientação a '
    'output, para avaliar a eficiência relativa de 22 grandes empresas do setor de petróleo e '
    'gás, identificar os benchmarks setoriais e propor metas de melhoria para as unidades ineficientes.'
))

# ════════════════════════════════════════════════════════════════════════════
# APRESENTAÇÃO DO PROBLEMA
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, 'APRESENTAÇÃO DO PROBLEMA')

subheading_inline(doc, 'Unidades de Decisão (DMUs). ',
    'A amostra é composta por 22 empresas do setor de petróleo e gás listadas em bolsas '
    'internacionais, abrangendo majors integradas, independentes de exploração, empresas de '
    'refino, serviços e midstream. As empresas selecionadas são megacorporações globais com '
    'operações em múltiplos países e dados reportados em dólares americanos, o que minimiza '
    'distorções de comparação internacional típicas de estudos com unidades puramente domésticas. '
    'A composição da amostra é apresentada na Tabela 1.'
)
make_table(doc,
    headers=['Ticker', 'Empresa', 'País', 'Segmento'],
    rows=[
        ('XOM','ExxonMobil','EUA','Integrada'),('CVX','Chevron','EUA','Integrada'),
        ('SHEL','Shell','Reino Unido','Integrada'),('BP','BP','Reino Unido','Integrada'),
        ('TTE','TotalEnergies','França','Integrada'),('COP','ConocoPhillips','EUA','Exploração'),
        ('PBR','Petrobras','Brasil','Integrada'),('EQNR','Equinor','Noruega','Integrada'),
        ('E','Eni','Itália','Integrada'),('OXY','Occidental Petroleum','EUA','Exploração'),
        ('DVN','Devon Energy','EUA','Exploração'),('VLO','Valero Energy','EUA','Refino'),
        ('PSX','Phillips 66','EUA','Refino'),('MPC','Marathon Petroleum','EUA','Refino'),
        ('SLB','SLB (Schlumberger)','EUA','Serviços'),('HAL','Halliburton','EUA','Serviços'),
        ('IMO','Imperial Oil','Canadá','Integrada'),('CNQ','Canadian Natural Resources','Canadá','Exploração'),
        ('SU','Suncor Energy','Canadá','Integrada'),('BKR','Baker Hughes','EUA','Serviços'),
        ('FTI','TechnipFMC','EUA/França','Serviços'),('WMB','Williams Companies','EUA','Midstream'),
    ],
    col_widths_cm=[2.0, 5.5, 4.0, 5.06],
    caption_label='Tabela 1. ',
    caption_text='Relação das DMUs analisadas, com ticker, empresa, país de origem e segmento de atuação.'
)

subheading_inline(doc, 'Variáveis. ',
    'A escolha de variáveis seguiu os critérios de disponibilidade, representatividade econômica '
    'e ausência de multicolinearidade severa. Foram definidos três inputs e dois outputs, '
    'conforme apresentado na Tabela 2.'
)
make_table(doc,
    headers=['Tipo', 'Variável', 'Descrição'],
    rows=[
        ('Input','Total Assets','Ativos totais — proxy para capital empregado'),
        ('Input','Operating Expenses','Despesas operacionais — proxy para custos'),
        ('Input','Employees','Nº de funcionários — proxy para capital humano'),
        ('Output','Total Revenue','Receita líquida total'),
        ('Output','EBITDA','Lucro antes de juros, impostos, depreciação e amortização'),
    ],
    col_widths_cm=[2.5, 4.5, 9.56],
    caption_label='Tabela 2. ',
    caption_text='Variáveis utilizadas no modelo DEA, com tipo (input/output) e descrição.'
)
body_para(doc, (
    'A regra prática do DEA recomenda que o número de DMUs seja pelo menos três vezes a soma '
    'de inputs e outputs: 3 × (3 + 2) = 15. Com 22 DMUs, a regra é atendida com margem de '
    '7 unidades. Os dados foram coletados via API do Yahoo Finance (biblioteca yfinance, Python), '
    'com valores expressos em dólares americanos. Cinco empresas foram removidas por dados '
    'incompletos (REP, MRO, HES, YPF, EC).'
))
subheading_inline(doc, 'Limitações. ',
    'A fonte automática (Yahoo Finance) pode apresentar inconsistências para empresas em bolsas '
    'não americanas. Os anos fiscais não são uniformes entre todas as empresas, podendo introduzir '
    'viés temporal. A Petrobras (PBR) foi identificada como outlier em três das cinco variáveis, '
    'podendo distorcer a fronteira eficiente. A amostra inclui empresas de segmentos distintos, '
    'o que reduz parcialmente a homogeneidade das DMUs.'
)

# ════════════════════════════════════════════════════════════════════════════
# ANÁLISE ENVOLTÓRIA DE DADOS
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, 'ANÁLISE ENVOLTÓRIA DE DADOS')

subheading_inline(doc, 'Introdução ao DEA. ',
    'A Análise Envoltória de Dados é uma técnica de programação matemática não paramétrica que '
    'avalia a eficiência relativa de unidades de decisão que compartilham o mesmo conjunto de '
    'inputs e outputs. Ao contrário de métodos paramétricos, o DEA não pressupõe uma forma '
    'funcional específica, construindo a fronteira eficiente empiricamente a partir das melhores '
    'práticas observadas na amostra. Uma DMU é eficiente quando nenhuma combinação linear das '
    'demais consegue produzir mais outputs com os mesmos ou menos inputs.'
)
subheading_inline(doc, 'Modelo CCR. ',
    'O modelo CCR (Charnes, Cooper e Rhodes, 1978) assume retornos constantes de escala (CRS). '
    'Com orientação a output, busca o fator máximo \u03b8 pelo qual os outputs da DMU avaliada '
    'podem ser expandidos mantendo os inputs fixos:'
)
add_equation(doc, 'max  \u03b8', '(1)')
add_equation(doc, 's.a.    \u03a3j \u03bbj xij \u2264 xio ,   \u2200 i', '(2)')
add_equation(doc, '         \u03a3j \u03bbj yrj \u2265 \u03b8 \u00b7 yro ,  \u2200 r', '(3)')
add_equation(doc, '         \u03bbj \u2265 0 ,   \u2200 j', '(4)')

body_para(doc, (
    'onde xij é o valor do input i da DMU j; yrj é o valor do output r da DMU j; \u03bbj são '
    'os pesos das DMUs de referência; e \u03b8 é o score de eficiência. Uma DMU é eficiente '
    'quando \u03b8 = 1,0; ineficiente quando \u03b8 > 1,0.'
))
subheading_inline(doc, 'Modelo BCC. ',
    'O modelo BCC (Banker, Charnes e Cooper, 1984) admite retornos variáveis de escala (VRS) '
    'por meio de uma restrição de convexidade adicional:'
)
add_equation(doc, 'max  \u03b8', '(5)')
add_equation(doc, 's.a.    \u03a3j \u03bbj xij \u2264 xio ,   \u2200 i', '(6)')
add_equation(doc, '         \u03a3j \u03bbj yrj \u2265 \u03b8 \u00b7 yro ,  \u2200 r', '(7)')
add_equation(doc, '         \u03bbj \u2265 0 ,   \u2200 j', '(8)')
add_equation(doc, '         \u03a3j \u03bbj = 1', '(9)')

body_para(doc, (
    'A restrição (9) garante que cada DMU seja comparada apenas com combinações convexas de '
    'outras DMUs de porte similar, eliminando o componente de ineficiência de escala.'
))
subheading_inline(doc, 'Eficiência de Escala. ',
    'A eficiência de escala (SE) é obtida pela razão entre os scores dos dois modelos:'
)
add_equation(doc, 'SE = \u03b8CCR / \u03b8BCC', '(10)')
body_para(doc, (
    'Quando SE = 1, a DMU opera na escala ótima (CRS). Quando SE > 1, há ineficiência de escala.'
))
subheading_inline(doc, 'Orientação escolhida. ',
    'Adotou-se a orientação a output pois as empresas do setor operam com estrutura de ativos, '
    'força de trabalho e custos operacionais relativamente fixos no curto prazo. A pergunta '
    'economicamente mais relevante é quanto mais receita e EBITDA cada empresa poderia gerar '
    'com os recursos que já possui, caso operasse sobre a fronteira eficiente.'
)

# ════════════════════════════════════════════════════════════════════════════
# RESULTADOS E DISCUSSÕES
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
section_heading(doc, 'RESULTADOS E DISCUSSÕES')

subheading_inline(doc, 'Análise Exploratória dos Dados. ',
    'A análise exploratória revelou elevada dispersão nas variáveis, reflexo da heterogeneidade '
    'de porte entre as 22 DMUs. Os boxplots das variáveis normalizadas (Figura 5) evidenciam a '
    'presença de valores atípicos, especialmente nos outputs. A análise pelo método IQR '
    'identificou a Petrobras (PBR) como outlier em três variáveis: Total Assets, Receita '
    'Líquida e EBITDA.'
)
add_figure(doc, os.path.join(FIGS2, 'boxplots.png'),
    'Figura 5. ', 'Boxplots das variáveis de input e output normalizadas pelo valor máximo da amostra.',
    width=Cm(14))

body_para(doc, (
    'A matriz de correlação (Figura 6) apresenta alta correlação positiva entre inputs e outputs, '
    'esperada pela relação entre porte e capacidade produtiva das empresas.'
))
add_figure(doc, os.path.join(FIGS2, 'correlation_heatmap.png'),
    'Figura 6. ', 'Matriz de correlação entre as variáveis de input e output.',
    width=Cm(11))

subheading_inline(doc, 'Scores de Eficiência. ',
    'Os resultados dos modelos CCR e BCC são apresentados na Tabela 3. No modelo CCR, 9 das '
    '22 DMUs (41%) atingiram a fronteira eficiente (\u03b8 = 1,0000), enquanto no BCC esse '
    'número sobe para 10 DMUs (45%). O score médio CCR é 1,0811 e o BCC é 1,0591.'
)
make_table(doc,
    headers=['DMU','Empresa','CCR','BCC','Ef. Escala','Efic. CCR','Efic. BCC'],
    rows=[
        ('XOM','ExxonMobil','1,1023','1,0000','1,1023','Não','Sim'),
        ('CVX','Chevron','1,1849','1,1281','1,0504','Não','Não'),
        ('SHEL','Shell','1,1356','1,0188','1,1146','Não','Não'),
        ('BP','BP','1,1522','1,0686','1,0782','Não','Não'),
        ('TTE','TotalEnergies','1,1495','1,0854','1,0591','Não','Não'),
        ('COP','ConocoPhillips','1,0587','1,0585','1,0002','Não','Não'),
        ('PBR','Petrobras','1,0000','1,0000','1,0000','Sim','Sim'),
        ('EQNR','Equinor','1,0000','1,0000','1,0000','Sim','Sim'),
        ('E','Eni','1,1942','1,1937','1,0004','Não','Não'),
        ('OXY','Occidental','1,1281','1,1253','1,0025','Não','Não'),
        ('DVN','Devon Energy','1,0000','1,0000','1,0000','Sim','Sim'),
        ('VLO','Valero','1,0000','1,0000','1,0000','Sim','Sim'),
        ('PSX','Phillips 66','1,0000','1,0000','1,0000','Sim','Sim'),
        ('MPC','Marathon Petro.','1,0000','1,0000','1,0000','Sim','Sim'),
        ('SLB','SLB','1,0976','1,0889','1,0080','Não','Não'),
        ('HAL','Halliburton','1,0504','1,0488','1,0015','Não','Não'),
        ('IMO','Imperial Oil','1,0000','1,0000','1,0000','Sim','Sim'),
        ('CNQ','Canadian Natural','1,2597','1,2438','1,0128','Não','Não'),
        ('SU','Suncor Energy','1,1513','1,1336','1,0156','Não','Não'),
        ('BKR','Baker Hughes','1,1199','1,1077','1,0110','Não','Não'),
        ('FTI','TechnipFMC','1,0000','1,0000','1,0000','Sim','Sim'),
        ('WMB','Williams Cos.','1,0000','1,0000','1,0000','Sim','Sim'),
    ],
    col_widths_cm=[1.5, 3.8, 1.8, 1.8, 2.2, 2.2, 2.26],
    caption_label='Tabela 3. ',
    caption_text='Scores de eficiência CCR, BCC e eficiência de escala para as 22 DMUs. Scores iguais a 1,0000 indicam DMU eficiente.'
)

body_para(doc, 'A distribuição dos scores é visualizada no histograma da Figura 1.')
add_figure(doc, os.path.join(FIGS4, 'histogram_efficiency.png'),
    'Figura 1. ', 'Distribuição dos scores de eficiência CCR e BCC. A linha vermelha tracejada indica a fronteira eficiente (\u03b8 = 1,0).',
    width=Cm(14))

body_para(doc, 'O ranking de eficiência por empresa é apresentado na Figura 2.')
add_figure(doc, os.path.join(FIGS4, 'ranking_efficiency.png'),
    'Figura 2. ', 'Ranking de eficiência das 22 DMUs pelos scores CCR e BCC.',
    width=Cm(14))

subheading_inline(doc, 'DMUs Eficientes. ',
    'As empresas que atingiram a fronteira eficiente em cada modelo são sintetizadas na Tabela 4.'
)
make_table(doc,
    headers=['Modelo','DMUs Eficientes (\u03b8 = 1,0000)'],
    rows=[
        ('CCR','PBR, EQNR, DVN, VLO, PSX, MPC, IMO, FTI, WMB (9 empresas)'),
        ('BCC','XOM, PBR, EQNR, DVN, VLO, PSX, MPC, IMO, FTI, WMB (10 empresas)'),
    ],
    col_widths_cm=[2.5, 14.06],
    caption_label='Tabela 4. ',
    caption_text='DMUs eficientes nos modelos CCR e BCC. XOM aparece eficiente exclusivamente no BCC, evidenciando ineficiência de escala.'
)

subheading_inline(doc, 'Benchmarks e Lambdas. ',
    'O scatter plot CCR vs BCC (Figura 3) e o heatmap de benchmarks (Figura 4) evidenciam '
    'a estrutura de referências do modelo.'
)
add_figure(doc, os.path.join(FIGS4, 'ccr_vs_bcc.png'),
    'Figura 3. ', 'Dispersão entre scores CCR e BCC por DMU. Pontos sobre a diagonal (y = x) indicam ausência de ineficiência de escala.',
    width=Cm(12))
add_figure(doc, os.path.join(FIGS4, 'benchmark_heatmap.png'),
    'Figura 4. ', 'Heatmap de intensidade dos pesos \u03bb (modelo CCR). Linhas: DMUs ineficientes; Colunas: DMUs eficientes utilizadas como benchmark.',
    width=Cm(14))

make_table(doc,
    headers=['DMU Ineficiente','Score CCR','Benchmarks (CCR)'],
    rows=[
        ('XOM','1,1023','PBR, EQNR, IMO'),('CVX','1,1849','PBR, EQNR, IMO'),
        ('SHEL','1,1356','EQNR, MPC, FTI'),('BP','1,1522','EQNR, FTI'),
        ('TTE','1,1495','PBR, EQNR'),('COP','1,0587','PBR, EQNR, DVN'),
        ('E','1,1942','PBR, EQNR'),('OXY','1,1281','PBR, WMB'),
        ('SLB','1,0976','PBR, EQNR'),('HAL','1,0504','EQNR, FTI'),
        ('CNQ','1,2597','PBR, EQNR, DVN'),('SU','1,1513','PBR, EQNR'),
        ('BKR','1,1199','PBR, EQNR'),
    ],
    col_widths_cm=[4.0, 3.0, 9.56],
    caption_label='Tabela 5. ',
    caption_text='Benchmarks das DMUs ineficientes no modelo CCR. EQNR é referência para 12 das 13 DMUs ineficientes; PBR, para 10.'
)
body_para(doc, (
    'A Equinor (EQNR) é o benchmark dominante, servindo de referência para 12 das 13 DMUs '
    'ineficientes no modelo CCR. Seu perfil — alta receita e EBITDA por ativo empregado, com '
    'base de funcionários enxuta (aproximadamente 22.000 colaboradores) — representa o padrão '
    'de excelência setorial identificado pelo DEA.'
))

subheading_inline(doc, 'Metas de Melhoria. ',
    'A orientação a output implica que as DMUs ineficientes deveriam expandir receita e EBITDA '
    'em (\u03b8 \u2212 1) \u00d7 100% mantendo os inputs atuais. As metas são apresentadas '
    'na Tabela 6 e visualizadas na Figura 7.'
)
make_table(doc,
    headers=['DMU','Score CCR','Receita Atual\n(USD bi)','Meta Receita\n(USD bi)','Aumento',
             'EBITDA Atual\n(USD bi)','Meta EBITDA\n(USD bi)','Aumento'],
    rows=[
        ('CNQ','1,2597','38,6','48,7','+26,0%','14,8','18,6','+26,0%'),
        ('E (Eni)','1,1942','84,5','101,0','+19,4%','12,0','14,3','+19,4%'),
        ('CVX','1,1849','186,0','220,4','+18,5%','37,9','44,9','+18,5%'),
    ],
    col_widths_cm=[1.8, 1.8, 2.0, 2.0, 1.6, 2.0, 2.0, 1.76] ,
    caption_label='Tabela 6. ',
    caption_text='Metas de melhoria para as três DMUs mais ineficientes no modelo CCR. Valores em bilhões de dólares americanos (USD bi).'
)
add_figure(doc, os.path.join(FIGS4, 'improvement_targets.png'),
    'Figura 7. ', 'Comparação entre outputs atuais e metas para as três DMUs mais ineficientes (CNQ, Eni e Chevron).',
    width=Cm(14))
body_para(doc, (
    'A Canadian Natural Resources (CNQ), com \u03b8 = 1,2597, é a mais ineficiente. Sua '
    'ineficiência é majoritariamente estrutural, associada ao modelo de negócio de oil sands, '
    'cujo custo de extração é superior ao do petróleo convencional explorado por seus benchmarks. '
    'A Eni (E), com \u03b8 = 1,1942, apresenta base de ativos e funcionários elevada em relação '
    'ao nível de receita. A Chevron (CVX), com \u03b8 = 1,1849, não extrai o máximo de receita '
    'e EBITDA de sua base de ativos, possivelmente por estratégia conservadora de capital.'
))

subheading_inline(doc, 'Eficiência de Escala. ',
    'A eficiência de escala de cada DMU é apresentada na Tabela 7. O caso mais relevante é o '
    'da ExxonMobil (XOM), com SE = 1,1023 — a maior da amostra. Seu score CCR cai para 1,0000 '
    'no BCC, evidenciando que toda a ineficiência observada no modelo CCR é de natureza escalar, '
    'não operacional.'
)
make_table(doc,
    headers=['DMU','Score CCR','Score BCC','Ef. Escala','Retorno de Escala'],
    rows=[
        ('XOM','1,1023','1,0000','1,1023','CRS'),('CVX','1,1849','1,1281','1,0504','CRS'),
        ('SHEL','1,1356','1,0188','1,1146','CRS'),('BP','1,1522','1,0686','1,0782','CRS'),
        ('TTE','1,1495','1,0854','1,0591','CRS'),('COP','1,0587','1,0585','1,0002','CRS'),
        ('PBR','1,0000','1,0000','1,0000','CRS'),('EQNR','1,0000','1,0000','1,0000','CRS'),
        ('E','1,1942','1,1937','1,0004','CRS'),('OXY','1,1281','1,1253','1,0025','CRS'),
        ('DVN','1,0000','1,0000','1,0000','CRS'),('VLO','1,0000','1,0000','1,0000','CRS'),
        ('PSX','1,0000','1,0000','1,0000','CRS'),('MPC','1,0000','1,0000','1,0000','CRS'),
        ('SLB','1,0976','1,0889','1,0080','CRS'),('HAL','1,0504','1,0488','1,0015','CRS'),
        ('IMO','1,0000','1,0000','1,0000','CRS'),('CNQ','1,2597','1,2438','1,0128','CRS'),
        ('SU','1,1513','1,1336','1,0156','CRS'),('BKR','1,1199','1,1077','1,0110','CRS'),
        ('FTI','1,0000','1,0000','1,0000','CRS'),('WMB','1,0000','1,0000','1,0000','CRS'),
    ],
    col_widths_cm=[2.0, 2.8, 2.8, 2.8, 6.16],
    caption_label='Tabela 7. ',
    caption_text='Eficiência de escala e tipo de retorno de escala para todas as DMUs.'
)

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSÕES
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
section_heading(doc, 'CONCLUSÕES')
body_para(doc, (
    'O presente trabalho aplicou os modelos DEA-CCR e DEA-BCC, com orientação a output, para '
    'avaliar a eficiência relativa de 22 grandes empresas do setor de petróleo e gás. '
    '41% das DMUs (9/22) são eficientes no CCR e 45% (10/22) no BCC, com scores médios de '
    '1,0811 e 1,0591, respectivamente. A Equinor (EQNR) e a Petrobras (PBR) são os benchmarks '
    'dominantes — referências para 12 e 10 DMUs ineficientes, respectivamente. A Canadian '
    'Natural Resources (CNQ) é a mais ineficiente (\u03b8 = 1,2597), com ineficiência '
    'majoritariamente estrutural ligada ao modelo de negócio de oil sands. A ExxonMobil (XOM) '
    'apresenta o caso mais didático de ineficiência de escala: operacionalmente eficiente no '
    'BCC, mas com SE = 1,1023 no CCR. As grandes majors europeias — Shell, BP, TotalEnergies '
    'e Eni — apresentam-se como ineficientes em ambos os modelos, com scores entre 1,02 e 1,19.'
))
body_para(doc, (
    'Para trabalhos futuros, recomenda-se a análise de janela temporal (window analysis) para '
    'capturar a evolução da eficiência ao longo dos ciclos de preço do petróleo, a inclusão '
    'de variáveis ambientais e de governança (ESG) como outputs, e a aplicação de modelos DEA '
    'com outputs indesejáveis (emissões de CO\u2082) para refletir as externalidades do setor.'
))

# ════════════════════════════════════════════════════════════════════════════
# REFERÊNCIAS BIBLIOGRÁFICAS
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, 'REFERÊNCIAS BIBLIOGRÁFICAS')

refs = [
    ('BANKER, R. D.; CHARNES, A.; COOPER, W. W. ',
     'Some models for estimating technical and scale inefficiencies in data envelopment analysis. ',
     'Management Science', ', v. 30, n. 9, p. 1078–1092, 1984.'),
    ('CHARNES, A.; COOPER, W. W.; RHODES, E. ',
     'Measuring the efficiency of decision making units. ',
     'European Journal of Operational Research', ', v. 2, n. 6, p. 429–444, 1978.'),
    ('COOPER, W. W.; SEIFORD, L. M.; TONE, K. ',
     'Data Envelopment Analysis: A Comprehensive Text with Models, Applications, References and DEA-Solver Software',
     '. 2. ed. New York: Springer, 2007.', ''),
    ('LINS, M. P. E.; MEZA, L. A. ',
     'Análise Envoltória de Dados e Perspectivas de Integração no Ambiente de Apoio à Decisão',
     '. Rio de Janeiro: COPPE/UFRJ, 2000.', ''),
    ('MITCHELL, J. ',
     'yfinance: Yahoo! Finance market data downloader',
     '. Disponível em: https://pypi.org/project/yfinance/. Acesso em: jun. 2025.', ''),
    ('MITCHELL, S. et al. ',
     'PuLP: A Linear Programming toolkit for Python',
     '. Disponível em: https://pypi.org/project/PuLP/. Acesso em: jun. 2025.', ''),
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.line_spacing = Pt(14)
    add_run(p, ref[0], size=12)
    add_run(p, ref[1], italic=True, size=12)
    add_run(p, ref[2], size=12)
    if len(ref) > 3:
        add_run(p, ref[3], size=12)

# ── Salvar ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'Relatório salvo em: {OUT}')
print(f'Tamanho: {os.path.getsize(OUT) / 1024:.0f} KB')
